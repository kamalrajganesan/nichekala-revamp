"""Generate the client project intake .docx template for Nichekala dev workflow."""
from pathlib import Path
from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


SHOTS_DIR = Path(__file__).resolve().parent / "screenshots"


BRAND = RGBColor(0xC8, 0x15, 0x7B)
GRAY = RGBColor(0x88, 0x88, 0x88)
DARK = RGBColor(0x22, 0x22, 0x22)


def add_h1(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(6)
    r = p.add_run(text)
    r.font.size = Pt(22)
    r.font.bold = True
    r.font.color.rgb = DARK


def add_h2(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(18)
    p.paragraph_format.space_after = Pt(4)
    r = p.add_run(text)
    r.font.size = Pt(13)
    r.font.bold = True
    r.font.color.rgb = BRAND


def add_label(doc, label, hint=""):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after = Pt(2)
    r = p.add_run(label)
    r.font.size = Pt(10)
    r.font.bold = True
    r.font.color.rgb = DARK
    if hint:
        h = p.add_run("  " + hint)
        h.font.size = Pt(9)
        h.font.italic = True
        h.font.color.rgb = GRAY


def add_placeholder(doc, placeholder):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(2)
    r = p.add_run(placeholder)
    r.font.size = Pt(10)
    r.font.italic = True
    r.font.color.rgb = GRAY


def add_blank_line(doc, lines=1):
    for _ in range(lines):
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(0)
        r = p.add_run(" ")
        r.font.size = Pt(10)


def add_note(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.left_indent = Cm(0.5)
    r = p.add_run("Note for dev: ")
    r.font.size = Pt(9)
    r.font.bold = True
    r.font.color.rgb = BRAND
    body = p.add_run(text)
    body.font.size = Pt(9)
    body.font.italic = True
    body.font.color.rgb = GRAY


def add_body(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(4)
    r = p.add_run(text)
    r.font.size = Pt(10)
    r.font.color.rgb = DARK


def add_screenshot(doc, filename, caption):
    img_path = SHOTS_DIR / filename
    if not img_path.exists():
        return
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(2)
    run = p.add_run()
    run.add_picture(str(img_path), width=Cm(16))
    cap = doc.add_paragraph()
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cap.paragraph_format.space_after = Pt(10)
    cr = cap.add_run(caption)
    cr.font.size = Pt(8)
    cr.font.italic = True
    cr.font.color.rgb = GRAY


def add_divider(doc):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(14)
    p.paragraph_format.space_after = Pt(8)
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "6")
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), "C8157B")
    pBdr.append(bottom)
    pPr.append(pBdr)


def build():
    doc = Document()

    # Tighten default margins
    for section in doc.sections:
        section.left_margin = Cm(2)
        section.right_margin = Cm(2)
        section.top_margin = Cm(1.8)
        section.bottom_margin = Cm(1.8)

    # Title block
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.LEFT
    r = title.add_run("NICHEKALA.")
    r.font.size = Pt(11)
    r.font.bold = True
    r.font.color.rgb = BRAND

    add_h1(doc, "Project Intake Form")
    sub = doc.add_paragraph()
    sr = sub.add_run(
        "Please fill in the details below for your project. "
        "Items marked * are required. Everything else helps us tell your story better, "
        "but skip anything that doesn't apply."
    )
    sr.font.size = Pt(10)
    sr.font.color.rgb = GRAY

    add_divider(doc)

    # ──────────────────────────────────────────────────────────────
    # SECTION 1 — Project Identity
    # ──────────────────────────────────────────────────────────────
    add_h2(doc, "1.  Project Identity")
    add_screenshot(doc, "01-identity.png", "↑  How it appears on the page — breadcrumbs, project title, scroll-down arrow.")

    add_label(doc, "Project name *", "(e.g. \"Inside Out Residence\", \"Coffee and Conversations\")")
    add_placeholder(doc, "Type project name here…")
    add_blank_line(doc)

    add_label(doc, "Tagline / subtitle", "(optional one-line subtitle that sits below the project name)")
    add_placeholder(doc, "e.g. \"A 1200 sqft canvas of light and openness\"")
    add_blank_line(doc)

    # ──────────────────────────────────────────────────────────────
    # SECTION 2 — Project Facts
    # ──────────────────────────────────────────────────────────────
    add_h2(doc, "2.  Project Facts")
    add_screenshot(doc, "02-facts.png", "↑  The Build Type / Location / Status row appears below the hero image.")

    add_label(doc, "Build type *", "(pick one: Residential / Commercial / Recreational / Heritage / Pocket Home / Other)")
    add_placeholder(doc, "e.g. Residential")
    add_blank_line(doc)

    add_label(doc, "Location *", "(City, State — e.g. \"Chennai, Tamil Nadu\" or \"Goa, India\")")
    add_placeholder(doc, "Type location here…")
    add_blank_line(doc)

    add_label(doc, "Status *", "(pick one: Proposed / Under Construction / Completed)")
    add_placeholder(doc, "e.g. Completed")
    add_blank_line(doc)

    add_label(doc, "Year", "(year of completion or proposal — e.g. \"2023\" or \"2022–2024\")")
    add_placeholder(doc, "e.g. 2024")
    add_blank_line(doc)

    add_label(doc, "Plot size / built-up area", "(optional — e.g. \"1200 sqft plot\", \"20×60 plot\", \"3 BHK / 2400 sqft built-up\")")
    add_placeholder(doc, "Type area details here…")
    add_blank_line(doc)

    # ──────────────────────────────────────────────────────────────
    # SECTION 3 — The Story (Narrative)
    # ──────────────────────────────────────────────────────────────
    add_h2(doc, "3.  The Story")
    add_screenshot(doc, "03-story.png", "↑  Your narrative paragraph appears here, to the right of the project facts.")

    add_label(doc, "One-line summary *", "(used on portfolio listing + meta description — max ~25 words)")
    add_placeholder(doc, "e.g. A 1200 sqft residence in Chennai opening out to the park, with double-height living and rear-facing balconies.")
    add_blank_line(doc)

    add_label(
        doc,
        "Project narrative *",
        "(1–3 paragraphs — the brief, the design idea, what makes it special. "
        "Write naturally, our team will edit for the site.)",
    )
    add_placeholder(doc, "Paragraph 1: The brief / client need…")
    add_blank_line(doc)
    add_placeholder(doc, "Paragraph 2: The design idea / approach…")
    add_blank_line(doc)
    add_placeholder(doc, "Paragraph 3: Outcome / what makes it special…")
    add_blank_line(doc)

    add_label(doc, "Key design highlights", "(optional — 3 to 5 short bullets, one feature each)")
    add_placeholder(doc, "•  Double-height living space")
    add_placeholder(doc, "•  Rear-facing balconies open to community park")
    add_placeholder(doc, "•  Compact ground-floor guest room")
    add_placeholder(doc, "•  …")
    add_blank_line(doc)

    # ──────────────────────────────────────────────────────────────
    # SECTION 4 — Visuals
    # ──────────────────────────────────────────────────────────────
    add_h2(doc, "4.  Visuals")
    add_screenshot(doc, "04-visuals.png", "↑  The hero image dominates the banner; gallery photos display in a 2-column grid below.")
    add_note(
        doc,
        "Hero image is the big banner at the top. Gallery images appear below the description. "
        "Recommended specs: hero ≥ 1920×1080 px (landscape), gallery ≥ 1600×1200 px, "
        "WebP or high-quality JPG preferred. File names: use lowercase + hyphens, no spaces.",
    )

    add_label(doc, "Hero image *", "(main banner photo — provide filename + short alt text)")
    add_placeholder(doc, "Filename:  e.g. inside-out-hero.webp")
    add_placeholder(doc, "Alt text:  e.g. \"Inside Out residence rear elevation\"")
    add_blank_line(doc)

    add_label(
        doc,
        "Gallery images *",
        "(list 3 to 8 photos in the order you want them shown — filename + alt text each)",
    )
    for i in range(1, 7):
        add_placeholder(doc, f"{i}.  Filename: __________________     Alt text: __________________")
    add_blank_line(doc)

    add_label(doc, "Social-share image", "(optional — the image shown when the page is shared on WhatsApp, Facebook, LinkedIn. Defaults to hero if blank. Ideal 1200×630.)")
    add_placeholder(doc, "Filename: __________________     Alt text: __________________")
    add_blank_line(doc)

    # ──────────────────────────────────────────────────────────────
    # SECTION 5 — Optional Extras
    # ──────────────────────────────────────────────────────────────
    add_h2(doc, "5.  Optional Extras")

    add_label(doc, "Client testimonial", "(if you want a quote on the page — name + role + quote)")
    add_placeholder(doc, "Quote: \"…\"")
    add_placeholder(doc, "Name & role: __________________")
    add_blank_line(doc)

    add_label(doc, "Design team / credits", "(who worked on it — lead architect, interior designer, contractor, photographer)")
    add_placeholder(doc, "e.g. Lead Architect: __________ ; Photographer: __________")
    add_blank_line(doc)

    add_label(doc, "Press / publications", "(any features in magazines, blogs, awards — title + link)")
    add_placeholder(doc, "e.g. ArchDaily — “Title of feature” — https://…")
    add_blank_line(doc)

    add_divider(doc)

    # ──────────────────────────────────────────────────────────────
    # APPENDIX — Dev-only
    # ──────────────────────────────────────────────────────────────
    add_h2(doc, "Appendix — For dev team only (do not send to client)")

    add_body(doc, "Use these defaults when the client leaves something blank.")
    add_blank_line(doc)

    add_label(doc, "Page slug", "(URL path — derived from project name, lowercase + hyphens)")
    add_placeholder(doc, "Example:  /projects/inside-out")
    add_blank_line(doc)

    add_label(doc, "Page <title>", "(SEO title tag — pattern below)")
    add_placeholder(doc, "Pattern:  {Project name} | {Build type} {Project type} in {City} | Nichekala")
    add_placeholder(doc, "Example:  Inside Out | Residential Architecture in Chennai | Nichekala")
    add_blank_line(doc)

    add_label(doc, "Meta description", "(derived from one-line summary; 150–160 chars)")
    add_placeholder(doc, "If summary > 160 chars, trim to first complete sentence under 160 chars.")
    add_blank_line(doc)

    add_label(doc, "Keywords", "(generate from build type + location + 2 to 3 design highlights)")
    add_placeholder(doc, "Pattern:  {Project name}, Nichekala, {build type} {city}, {highlight 1}, {highlight 2}, …")
    add_blank_line(doc)

    add_label(doc, "JSON-LD CreativeWork block", "(populate from sections 1, 2, 3, 4 — see existing projects/Inside-out.html as reference)")
    add_blank_line(doc)

    add_label(doc, "Prev / Next project links", "(dev decides — usually previous / next entry in projects/ folder by recency)")
    add_blank_line(doc)

    add_label(doc, "Sitemap + llms.txt + llms-full.txt", "(add the new URL with priority 0.7)")
    add_blank_line(doc)

    add_label(doc, "Internal QA checklist", "")
    add_placeholder(doc, "☐  Page loads at /projects/{slug} (clean URL via .htaccess)")
    add_placeholder(doc, "☐  All gallery images have alt text")
    add_placeholder(doc, "☐  Fancybox zoom works on every image")
    add_placeholder(doc, "☐  Prev/Next/All-projects links wired up")
    add_placeholder(doc, "☐  JSON-LD validates (search.google.com/test/rich-results)")
    add_placeholder(doc, "☐  Page added to sitemap.xml, llms.txt, llms-full.txt")
    add_placeholder(doc, "☐  Card added to portfolio.html grid (with correct data-category)")
    add_placeholder(doc, "☐  OG image displays correctly when shared (LinkedIn Post Inspector)")

    add_divider(doc)

    footer = doc.add_paragraph()
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    fr = footer.add_run("Nichekala Architecture and Design Studio · nichekala.in")
    fr.font.size = Pt(8)
    fr.font.color.rgb = GRAY

    out_path = "docs/client-project-intake-template.docx"
    doc.save(out_path)
    print(f"Saved: {out_path}")


if __name__ == "__main__":
    build()
